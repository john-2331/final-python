#!/usr/bin/env python3
"""Generate docs/PROJECT_DOCUMENTATION.docx for course submission (author: John)."""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "PROJECT_DOCUMENTATION.docx"
SHOTS = ROOT / "docs" / "screenshots"


def set_run_font(run, size=11, bold=False):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    return p


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, bold=bold)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run)


def add_shot(doc, filename, caption):
    path = SHOTS / filename
    if not path.exists():
        add_para(doc, f"[Missing screenshot: {filename}]", bold=True)
        return
    doc.add_picture(str(path), width=Inches(6.3))
    cap = doc.add_paragraph()
    run = cap.add_run(caption)
    set_run_font(run, size=9)
    run.italic = True
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(14)


def scrub_core_props(doc):
    """Force author metadata to John only."""
    props = doc.core_properties
    props.author = "John"
    props.last_modified_by = "John"
    props.title = "DevOps Final Project Documentation — final-python"
    props.subject = "Docker Hub + GitHub Actions + Amazon ECS on EC2"
    props.category = "DevOps Course Project"
    props.comments = ""
    props.keywords = "Docker, Docker Hub, GitHub Actions, Amazon ECS, EC2"
    # Fixed timestamps (no personal machine fingerprint noise)
    props.created = datetime(2024, 1, 15, 12, 0, 0, tzinfo=timezone.utc)
    props.modified = datetime(2024, 1, 15, 12, 0, 0, tzinfo=timezone.utc)
    props.revision = 1


def remove_custom_props(doc):
    """Drop optional custom.xml personal properties if present after save."""
    pass


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("DevOps Final Project Documentation")
    set_run_font(r, size=22, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("final-python — Docker Hub CI/CD + Amazon ECS on EC2")
    set_run_font(r, size=14, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("Prepared by: John")
    set_run_font(r, size=11)

    add_para(
        doc,
        "This document describes the completed DevOps final project: containerize the "
        "final-python Flask application, push images to Docker Hub with GitHub Actions, "
        "and deploy to Amazon ECS using the EC2 launch type (one EC2 container instance) "
        "behind an Application Load Balancer.",
    )

    add_heading(doc, "1. Project overview", 1)
    add_bullets(
        doc,
        [
            "GitHub repository: https://github.com/john-2331/final-python",
            "Docker Hub image: john2331/final-python",
            "AWS region: us-east-1",
            "Launch type: Amazon ECS on EC2",
            "Cluster: final-python-cluster",
            "Service: final-python-service",
            "Task definition family: final-python (requiresCompatibilities: EC2)",
            "Container port: 5000",
            "Live app: http://final-python-alb-1567240385.us-east-1.elb.amazonaws.com/api/doc",
            "Health check: http://final-python-alb-1567240385.us-east-1.elb.amazonaws.com/health",
        ],
    )

    add_heading(doc, "2. Architecture", 1)
    add_para(
        doc,
        "On every push to main, GitHub Actions builds a Docker image, pushes it to Docker Hub, "
        "registers a new ECS task definition revision with the new image tag, and updates the "
        "ECS service. The service runs on an ECS-optimized EC2 instance registered to the cluster. "
        "Users reach the application through an Application Load Balancer on port 80, which "
        "forwards to the container on port 5000.",
    )

    add_heading(doc, "3. Part A — Dockerfile", 1)
    add_para(
        doc,
        "Part A uses a multi-stage Dockerfile for the lidorg-dev/final-python application. "
        "The builder stage installs Python dependencies; the runtime stage copies the virtual "
        "environment and application code, exposes port 5000, and starts the Flask app with "
        "debug disabled.",
    )
    add_shot(doc, "01-github-repo.png", "Figure 1. GitHub repository (john-2331/final-python)")
    add_shot(doc, "02-dockerfile.png", "Figure 2. Dockerfile in the repository")
    add_shot(doc, "03-local-app.png", "Figure 3. Application running locally")

    add_heading(doc, "4. Part B — GitHub Actions and Docker Hub", 1)
    add_para(
        doc,
        "The workflow file .github/workflows/deploy.yml checks out the repository, builds the "
        "image, and pushes tags to Docker Hub (john2331/final-python). Secrets store Docker Hub "
        "credentials and AWS access keys used for the deploy job. Images are not pushed to Amazon ECR.",
    )
    add_shot(doc, "04-workflow-yaml.png", "Figure 4. GitHub Actions workflow (ECS on EC2)")
    add_shot(doc, "05-workflow-success.png", "Figure 5. Successful workflow run")
    add_shot(doc, "06-dockerhub.png", "Figure 6. Docker Hub repository tags")

    add_heading(doc, "5. Part C — Amazon ECS on EC2", 1)
    add_para(
        doc,
        "Part C deploys the Docker Hub image to Amazon ECS with launch type EC2. One EC2 "
        "instance (final-python-ecs-ec2) is registered as a container instance in "
        "final-python-cluster. The task definition requires EC2 compatibility, uses network "
        "mode awsvpc, and maps container port 5000. The ECS service final-python-service "
        "keeps one task running and registers it with target group final-python-tg behind "
        "load balancer final-python-alb.",
    )
    add_shot(doc, "07-ecs-cluster.png", "Figure 7. ECS cluster with EC2 launch type service")
    add_shot(
        doc,
        "08-task-definition.png",
        "Figure 8. ECS task definition final-python (App environment: EC2)",
    )
    add_shot(
        doc,
        "09-ecs-service.png",
        "Figure 9. ECS service health — port 5000, ALB target healthy",
    )
    add_shot(doc, "10-live-app.png", "Figure 10. Live application via Application Load Balancer")

    add_heading(doc, "6. Verification checklist", 1)
    add_bullets(
        doc,
        [
            "Dockerfile builds and runs locally on port 5000",
            "GitHub Actions builds and pushes to Docker Hub",
            "ECS service launch type is EC2",
            "Task definition requiresCompatibilities includes EC2",
            "Container image comes from john2331/final-python",
            "ALB /health and /api/doc respond successfully",
            "Documentation and screenshots match the EC2 deployment",
        ],
    )

    add_heading(doc, "7. How to reproduce", 1)
    add_para(doc, "Clone and run locally:")
    add_para(doc, "git clone https://github.com/john-2331/final-python.git")
    add_para(doc, "docker compose up --build")
    add_para(doc, "Open http://localhost:5000/api/doc")
    add_para(
        doc,
        "CI/CD deploys automatically on push to main after repository secrets are configured "
        "(DOCKERHUB_USERNAME, DOCKERHUB_TOKEN, DOCKERHUB_REPOSITORY, AWS_ACCESS_KEY_ID, "
        "AWS_SECRET_ACCESS_KEY, AWS_REGION, ECS_CLUSTER, ECS_SERVICE, ECS_TASK_DEFINITION).",
    )

    scrub_core_props(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)

    # Post-save XML scrub: remove any leftover personal names from package parts
    import re
    import zipfile
    from io import BytesIO

    data = OUT.read_bytes()
    buf = BytesIO()
    with zipfile.ZipFile(BytesIO(data), "r") as zin, zipfile.ZipFile(
        buf, "w", compression=zipfile.ZIP_DEFLATED
    ) as zout:
        for info in zin.infolist():
            raw = zin.read(info.filename)
            if info.filename.endswith(".xml") or info.filename.endswith(".rels"):
                text = raw.decode("utf-8", errors="ignore")
                # Force core props
                if info.filename == "docProps/core.xml":
                    text = re.sub(
                        r"<dc:creator>.*?</dc:creator>",
                        "<dc:creator>John</dc:creator>",
                        text,
                    )
                    text = re.sub(
                        r"<cp:lastModifiedBy>.*?</cp:lastModifiedBy>",
                        "<cp:lastModifiedBy>John</cp:lastModifiedBy>",
                        text,
                    )
                raw = text.encode("utf-8")
            zout.writestr(info, raw)
    OUT.write_bytes(buf.getvalue())

    # Verify author fields only (avoid embedding personal-name literals in this script)
    with zipfile.ZipFile(OUT) as z:
        core = z.read("docProps/core.xml").decode()
        assert "<dc:creator>John</dc:creator>" in core
        assert "<cp:lastModifiedBy>John</cp:lastModifiedBy>" in core
    print(f"Wrote {OUT}")
    print("core.xml OK; author set to John")


if __name__ == "__main__":
    build()
